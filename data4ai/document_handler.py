"""Document file operations for Data4AI - Extract and process PDF, DOCX, MD, and TXT files."""

import logging
import re
from pathlib import Path
from typing import Any, Optional, Union

from data4ai.exceptions import ValidationError

logger = logging.getLogger("data4ai")

# Check for optional dependencies
try:
    import pypdf
    PYPDF_AVAILABLE = True
except ImportError:
    PYPDF_AVAILABLE = False

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import markdown
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False


class DocumentHandler:
    """Handle document file operations for dataset generation."""

    @staticmethod
    def detect_document_type(file_path: Path) -> str:
        """Detect document type from file extension.
        
        Args:
            file_path: Path to document file
            
        Returns:
            Document type (pdf, docx, md, txt)
            
        Raises:
            ValidationError: If file type is not supported
        """
        suffix = file_path.suffix.lower()
        
        if suffix == ".pdf":
            return "pdf"
        elif suffix in [".docx", ".doc"]:
            return "docx"
        elif suffix in [".md", ".markdown"]:
            return "md"
        elif suffix in [".txt", ".text"]:
            return "txt"
        else:
            raise ValidationError(
                f"Unsupported document type: {suffix}. "
                "Supported types: .pdf, .docx, .md, .txt"
            )

    @staticmethod
    def extract_text(file_path: Path, use_advanced: bool = False) -> str:
        """Extract text from document file.
        
        Args:
            file_path: Path to document file
            use_advanced: Use advanced extraction (pdfplumber for PDFs)
            
        Returns:
            Extracted text content
            
        Raises:
            ValidationError: If extraction fails or dependencies missing
        """
        if not file_path.exists():
            raise FileNotFoundError(f"Document not found: {file_path}")
            
        doc_type = DocumentHandler.detect_document_type(file_path)
        
        if doc_type == "pdf":
            return DocumentHandler._extract_pdf_text(file_path, use_advanced)
        elif doc_type == "docx":
            return DocumentHandler._extract_docx_text(file_path)
        elif doc_type == "md":
            return DocumentHandler._extract_markdown_text(file_path)
        elif doc_type == "txt":
            return DocumentHandler._extract_txt_text(file_path)
        else:
            raise ValidationError(f"Unsupported document type: {doc_type}")

    @staticmethod
    def _extract_pdf_text(file_path: Path, use_advanced: bool = False) -> str:
        """Extract text from PDF file.
        
        Args:
            file_path: Path to PDF file
            use_advanced: Use pdfplumber for better extraction
            
        Returns:
            Extracted text
        """
        if use_advanced and PDFPLUMBER_AVAILABLE:
            logger.info("Using pdfplumber for advanced PDF extraction")
            try:
                import pdfplumber
                text_parts = []
                
                with pdfplumber.open(file_path) as pdf:
                    for page_num, page in enumerate(pdf.pages, 1):
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(f"--- Page {page_num} ---\n{page_text}")
                            
                return "\n\n".join(text_parts)
            except Exception as e:
                logger.warning(f"pdfplumber extraction failed: {e}, falling back to pypdf")
                # Fall back to pypdf
                
        if not PYPDF_AVAILABLE:
            raise ValidationError(
                "PDF support not available. Please install with: "
                "pip install data4ai[docs] or pip install pypdf"
            )
            
        try:
            import pypdf
            reader = pypdf.PdfReader(file_path)
            text_parts = []
            
            for page_num, page in enumerate(reader.pages, 1):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(f"--- Page {page_num} ---\n{page_text}")
                    
            return "\n\n".join(text_parts)
        except Exception as e:
            raise ValidationError(f"Failed to extract PDF text: {str(e)}") from e

    @staticmethod
    def _extract_docx_text(file_path: Path) -> str:
        """Extract text from DOCX file.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text
        """
        if not DOCX_AVAILABLE:
            raise ValidationError(
                "DOCX support not available. Please install with: "
                "pip install data4ai[docs] or pip install python-docx"
            )
            
        try:
            doc = Document(file_path)
            paragraphs = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
                    
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        paragraphs.append(row_text)
                        
            return "\n\n".join(paragraphs)
        except Exception as e:
            raise ValidationError(f"Failed to extract DOCX text: {str(e)}") from e

    @staticmethod
    def _extract_markdown_text(file_path: Path) -> str:
        """Extract text from Markdown file.
        
        Args:
            file_path: Path to Markdown file
            
        Returns:
            Plain text (with markdown stripped if library available)
        """
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
                
            if MARKDOWN_AVAILABLE:
                # Convert markdown to HTML then strip tags
                import markdown
                from html.parser import HTMLParser
                
                class HTMLStripper(HTMLParser):
                    def __init__(self):
                        super().__init__()
                        self.reset()
                        self.strict = False
                        self.convert_charrefs = True
                        self.text = []
                        
                    def handle_data(self, data):
                        self.text.append(data)
                        
                    def get_text(self):
                        return "".join(self.text)
                
                html = markdown.markdown(content)
                stripper = HTMLStripper()
                stripper.feed(html)
                return stripper.get_text()
            else:
                # Basic markdown stripping
                # Remove code blocks
                content = re.sub(r"```.*?```", "", content, flags=re.DOTALL)
                content = re.sub(r"`[^`]+`", "", content)
                # Remove headers
                content = re.sub(r"^#+\s+", "", content, flags=re.MULTILINE)
                # Remove emphasis
                content = re.sub(r"[*_]{1,2}([^*_]+)[*_]{1,2}", r"\1", content)
                # Remove links
                content = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", content)
                # Remove images
                content = re.sub(r"!\[([^\]]*)\]\([^)]+\)", "", content)
                
                return content
        except Exception as e:
            raise ValidationError(f"Failed to extract Markdown text: {str(e)}") from e

    @staticmethod
    def _extract_txt_text(file_path: Path) -> str:
        """Extract text from TXT file.
        
        Args:
            file_path: Path to TXT file
            
        Returns:
            File content
        """
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, "r", encoding="latin-1") as f:
                    return f.read()
            except Exception as e:
                raise ValidationError(f"Failed to extract text: {str(e)}") from e
        except Exception as e:
            raise ValidationError(f"Failed to extract text: {str(e)}") from e

    @staticmethod
    def extract_chunks(
        file_path: Path,
        chunk_size: int = 1000,
        overlap: int = 200,
        use_advanced: bool = False
    ) -> list[dict[str, Any]]:
        """Extract text in chunks for processing.
        
        Args:
            file_path: Path to document file
            chunk_size: Size of each chunk in characters
            overlap: Overlap between chunks for context
            use_advanced: Use advanced extraction
            
        Returns:
            List of chunks with metadata
        """
        text = DocumentHandler.extract_text(file_path, use_advanced)
        
        if not text:
            return []
            
        chunks = []
        start = 0
        chunk_id = 0
        
        while start < len(text):
            end = min(start + chunk_size, len(text))
            
            # Try to find a good break point (sentence end)
            if end < len(text):
                # Look for sentence endings
                for sep in [".\n", ". ", "!\n", "! ", "?\n", "? "]:
                    last_sep = text.rfind(sep, start, end)
                    if last_sep != -1:
                        end = last_sep + len(sep)
                        break
                        
            chunk_text = text[start:end].strip()
            
            if chunk_text:
                chunks.append({
                    "id": chunk_id,
                    "text": chunk_text,
                    "start": start,
                    "end": end,
                    "source": file_path.name,
                })
                chunk_id += 1
                
            # Move start position with overlap
            start = end - overlap if end < len(text) else end
            
        logger.info(f"Extracted {len(chunks)} chunks from {file_path.name}")
        return chunks

    @staticmethod
    def is_supported_file(file_path: Path) -> bool:
        """Check if file is a supported document type.
        
        Args:
            file_path: Path to check
            
        Returns:
            True if file is supported
        """
        try:
            DocumentHandler.detect_document_type(file_path)
            return True
        except ValidationError:
            return False

    @staticmethod
    def scan_folder(
        folder_path: Path,
        recursive: bool = True,
        file_types: Optional[list[str]] = None
    ) -> list[Path]:
        """Scan folder for supported documents.
        
        Args:
            folder_path: Path to folder
            recursive: Whether to scan subfolders
            file_types: Specific file types to include (pdf, docx, md, txt)
            
        Returns:
            List of document paths
        """
        if not folder_path.exists():
            raise FileNotFoundError(f"Folder not found: {folder_path}")
            
        if not folder_path.is_dir():
            raise ValidationError(f"Path is not a directory: {folder_path}")
            
        # Default to all supported types
        if file_types is None:
            file_types = ["pdf", "docx", "doc", "md", "markdown", "txt", "text"]
        else:
            # Normalize file types
            file_types = [ft.lower() for ft in file_types]
            
        documents = []
        
        # Define supported extensions
        extensions = []
        if "pdf" in file_types:
            extensions.append("*.pdf")
        if "docx" in file_types or "doc" in file_types:
            extensions.extend(["*.docx", "*.doc"])
        if "md" in file_types or "markdown" in file_types:
            extensions.extend(["*.md", "*.markdown"])
        if "txt" in file_types or "text" in file_types:
            extensions.extend(["*.txt", "*.text"])
            
        # Scan for files
        for ext in extensions:
            if recursive:
                documents.extend(folder_path.rglob(ext))
            else:
                documents.extend(folder_path.glob(ext))
                
        # Filter out hidden files and ensure unique paths
        documents = list(set([
            doc for doc in documents 
            if not doc.name.startswith(".") and doc.is_file()
        ]))
        
        # Sort for consistent ordering
        documents.sort()
        
        logger.info(f"Found {len(documents)} documents in {folder_path}")
        return documents

    @staticmethod
    def extract_from_multiple(
        file_paths: list[Path],
        use_advanced: bool = False,
        combine: bool = True
    ) -> Union[str, dict[str, str]]:
        """Extract text from multiple documents.
        
        Args:
            file_paths: List of document paths
            use_advanced: Use advanced extraction
            combine: Whether to combine into single text or return dict
            
        Returns:
            Combined text string or dict mapping paths to text
        """
        texts = {}
        
        for file_path in file_paths:
            try:
                text = DocumentHandler.extract_text(file_path, use_advanced)
                if text:
                    texts[str(file_path)] = text
            except Exception as e:
                logger.warning(f"Failed to extract from {file_path}: {e}")
                continue
                
        if combine:
            # Combine all texts with document separators
            combined_parts = []
            for path, text in texts.items():
                doc_name = Path(path).name
                combined_parts.append(f"=== Document: {doc_name} ===\n\n{text}")
            return "\n\n".join(combined_parts)
        else:
            return texts

    @staticmethod
    def extract_chunks_from_multiple(
        file_paths: list[Path],
        chunk_size: int = 1000,
        overlap: int = 200,
        use_advanced: bool = False
    ) -> list[dict[str, Any]]:
        """Extract chunks from multiple documents.
        
        Args:
            file_paths: List of document paths
            chunk_size: Size of each chunk
            overlap: Overlap between chunks
            use_advanced: Use advanced extraction
            
        Returns:
            List of all chunks with metadata
        """
        all_chunks = []
        
        for file_path in file_paths:
            try:
                chunks = DocumentHandler.extract_chunks(
                    file_path,
                    chunk_size=chunk_size,
                    overlap=overlap,
                    use_advanced=use_advanced
                )
                # Add file path to chunk metadata
                for chunk in chunks:
                    chunk["file_path"] = str(file_path)
                all_chunks.extend(chunks)
            except Exception as e:
                logger.warning(f"Failed to extract chunks from {file_path}: {e}")
                continue
                
        logger.info(f"Extracted {len(all_chunks)} total chunks from {len(file_paths)} documents")
        return all_chunks

    @staticmethod
    def prepare_for_generation(
        input_path: Union[Path, list[Path]],
        extraction_type: str = "qa",
        chunk_size: int = 1000,
        use_advanced: bool = False,
        recursive: bool = True
    ) -> dict[str, Any]:
        """Prepare document(s) for dataset generation.
        
        Args:
            input_path: Path to document file, folder, or list of paths
            extraction_type: Type of extraction (qa, summary, instruction)
            chunk_size: Size of chunks for processing
            use_advanced: Use advanced extraction
            recursive: For folders, whether to scan recursively
            
        Returns:
            Prepared data for generation
        """
        # Handle different input types
        if isinstance(input_path, list):
            # List of paths provided
            file_paths = input_path
            input_type = "multiple"
        elif isinstance(input_path, Path):
            if input_path.is_dir():
                # Folder provided - scan for documents
                file_paths = DocumentHandler.scan_folder(input_path, recursive=recursive)
                input_type = "folder"
                if not file_paths:
                    raise ValidationError(f"No supported documents found in {input_path}")
            else:
                # Single file provided
                file_paths = [input_path]
                input_type = "file"
        else:
            # Convert string to Path
            input_path = Path(input_path)
            return DocumentHandler.prepare_for_generation(
                input_path, extraction_type, chunk_size, use_advanced, recursive
            )
        
        # Extract chunks from all documents
        if len(file_paths) == 1:
            # Single document
            doc_type = DocumentHandler.detect_document_type(file_paths[0])
            chunks = DocumentHandler.extract_chunks(
                file_paths[0], chunk_size=chunk_size, use_advanced=use_advanced
            )
            
            return {
                "document_type": doc_type,
                "document_name": file_paths[0].name,
                "extraction_type": extraction_type,
                "chunks": chunks,
                "total_chunks": len(chunks),
                "total_documents": 1,
                "input_type": input_type,
                "metadata": {
                    "chunk_size": chunk_size,
                    "file_size": file_paths[0].stat().st_size,
                    "advanced_extraction": use_advanced,
                }
            }
        else:
            # Multiple documents
            chunks = DocumentHandler.extract_chunks_from_multiple(
                file_paths, chunk_size=chunk_size, use_advanced=use_advanced
            )
            
            # Get document types and sizes
            doc_types = set()
            total_size = 0
            for fp in file_paths:
                try:
                    doc_types.add(DocumentHandler.detect_document_type(fp))
                    total_size += fp.stat().st_size
                except:
                    pass
                    
            return {
                "document_type": "mixed" if len(doc_types) > 1 else list(doc_types)[0],
                "document_names": [fp.name for fp in file_paths],
                "extraction_type": extraction_type,
                "chunks": chunks,
                "total_chunks": len(chunks),
                "total_documents": len(file_paths),
                "input_type": input_type,
                "metadata": {
                    "chunk_size": chunk_size,
                    "total_file_size": total_size,
                    "advanced_extraction": use_advanced,
                    "document_types": list(doc_types),
                }
            }